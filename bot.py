import base64
import calendar
import json
import logging
import os
import re
import urllib.error
import urllib.request
import uuid
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
from openai import OpenAI
from telegram import InlineKeyboardButton, InlineKeyboardMarkup, KeyboardButton, ReplyKeyboardMarkup, Update
from telegram.ext import Application, CallbackQueryHandler, CommandHandler, ContextTypes, MessageHandler, filters

load_dotenv()
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(name)s: %(message)s")
logger = logging.getLogger(__name__)
logging.getLogger("httpx").setLevel(logging.WARNING)

TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
AI_PROVIDER = os.getenv("AI_PROVIDER", "openai").lower()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-5.2")
YANDEX_API_KEY = os.getenv("YANDEX_API_KEY")
YANDEX_FOLDER_ID = os.getenv("YANDEX_FOLDER_ID")
YANDEX_GPT_MODEL = os.getenv("YANDEX_GPT_MODEL", "yandexgpt-lite/latest")
YANDEX_OCR_MODEL = os.getenv("YANDEX_OCR_MODEL", "page")
CLEAN_CHAT = os.getenv("CLEAN_CHAT", "true").lower() == "true"
ADMIN_IDS = {admin_id.strip() for admin_id in os.getenv("ADMIN_IDS", "").split(",") if admin_id.strip()}
YOOKASSA_SHOP_ID = os.getenv("YOOKASSA_SHOP_ID", "")
YOOKASSA_SECRET_KEY = os.getenv("YOOKASSA_SECRET_KEY", "")
YOOKASSA_RETURN_URL = os.getenv("YOOKASSA_RETURN_URL", "https://t.me/")
CONTRACT_PRICE_RUB = os.getenv("CONTRACT_PRICE_RUB", "300")
PAYMENTS_ENABLED = os.getenv("PAYMENTS_ENABLED", "true").lower() == "true"
FREE_CONTRACT_LIMIT = int(os.getenv("FREE_CONTRACT_LIMIT", "2"))
DATA_DIR = Path(os.getenv("DATA_DIR", "."))
TEST_MODE_NOTICE = (
    "Бот работает в тестовом режиме. Возможны ошибки распознавания, поэтому перед подписанием "
    "обязательно проверь все данные в договоре вручную.\n\n"
    f"В тестовом режиме доступно {FREE_CONTRACT_LIMIT} генерации договора."
)

DATA_FILE = DATA_DIR / "deals.json"
USERS_FILE = DATA_DIR / "users.json"
OUTPUT_DIR = Path("contracts")
PHOTO_DIR = Path("tmp_photos")
TEMPLATE_FILE = Path("template.docx")

STAGES = [
    ("seller_passport", "Пришли фото паспорта продавца"),
    ("seller_registration", "Пришли фото прописки продавца"),
    ("buyer_passport", "Пришли фото паспорта покупателя"),
    ("buyer_registration", "Пришли фото прописки покупателя"),
    ("vehicle_sts_front", "Пришли фото СТС: первая сторона"),
    ("vehicle_sts_back", "Пришли фото СТС: вторая сторона"),
    ("vehicle_pts", "Пришли фото ПТС"),
]

EMPTY_DEAL = {
    "seller": {
        "full_name": "",
        "birth_date": "",
        "passport_series_number": "",
        "passport_issued_by": "",
        "passport_issue_date": "",
        "department_code": "",
        "registration_address": "",
    },
    "buyer": {
        "full_name": "",
        "birth_date": "",
        "passport_series_number": "",
        "passport_issued_by": "",
        "passport_issue_date": "",
        "department_code": "",
        "registration_address": "",
    },
    "vehicle": {
        "make_model": "",
        "vin": "",
        "vehicle_type": "",
        "year": "",
        "color": "",
        "license_plate": "",
        "pts": "",
        "pts_issued_by": "",
        "pts_issue_date": "",
        "sts": "",
        "sts_issued_by": "",
        "sts_issue_date": "",
        "body_number": "",
        "engine_number": "",
        "engine_power": "",
        "engine_volume": "",
        "chassis_number": "",
        "mileage": "",
    },
    "deal": {
        "city": "",
        "date": datetime.now().strftime("%d.%m.%Y"),
        "price": "",
    },
}

FIELD_LABELS = {
    "deal.city": "Город",
    "deal.date": "Дата",
    "deal.price": "Цена",
    "seller.full_name": "Продавец",
    "seller.birth_date": "Дата рождения продавца",
    "seller.passport_series_number": "Паспорт продавца",
    "seller.passport_issued_by": "Кем выдан паспорт продавца",
    "seller.passport_issue_date": "Дата выдачи паспорта продавца",
    "seller.department_code": "Код подразделения продавца",
    "seller.registration_address": "Прописка продавца",
    "buyer.full_name": "Покупатель",
    "buyer.birth_date": "Дата рождения покупателя",
    "buyer.passport_series_number": "Паспорт покупателя",
    "buyer.passport_issued_by": "Кем выдан паспорт покупателя",
    "buyer.passport_issue_date": "Дата выдачи паспорта покупателя",
    "buyer.department_code": "Код подразделения покупателя",
    "buyer.registration_address": "Прописка покупателя",
    "vehicle.make_model": "Автомобиль",
    "vehicle.vin": "VIN",
    "vehicle.vehicle_type": "Тип ТС",
    "vehicle.year": "Год",
    "vehicle.color": "Цвет",
    "vehicle.license_plate": "Госномер",
    "vehicle.pts": "ПТС",
    "vehicle.pts_issued_by": "Кем выдан ПТС",
    "vehicle.pts_issue_date": "Дата выдачи ПТС",
    "vehicle.sts": "СТС",
    "vehicle.sts_issued_by": "Кем выдан СТС",
    "vehicle.sts_issue_date": "Дата выдачи СТС",
    "vehicle.body_number": "Кузов",
    "vehicle.engine_number": "Двигатель",
    "vehicle.engine_power": "Мощность двигателя",
    "vehicle.engine_volume": "Объем двигателя",
    "vehicle.chassis_number": "Шасси/рама",
    "vehicle.mileage": "Пробег",
}

FIELD_ALIASES = {
    "город": "deal.city",
    "дата": "deal.date",
    "цена": "deal.price",
    "стоимость": "deal.price",
    "продавец": "seller.full_name",
    "фио продавца": "seller.full_name",
    "дата рождения продавца": "seller.birth_date",
    "паспорт продавца": "seller.passport_series_number",
    "паспорт серия номер продавца": "seller.passport_series_number",
    "кем выдан продавца": "seller.passport_issued_by",
    "дата выдачи продавца": "seller.passport_issue_date",
    "код подразделения продавца": "seller.department_code",
    "прописка продавца": "seller.registration_address",
    "адрес продавца": "seller.registration_address",
    "покупатель": "buyer.full_name",
    "фио покупателя": "buyer.full_name",
    "дата рождения покупателя": "buyer.birth_date",
    "паспорт покупателя": "buyer.passport_series_number",
    "паспорт серия номер покупателя": "buyer.passport_series_number",
    "паспорт серия номер": "buyer.passport_series_number",
    "паспорт серия номенр": "buyer.passport_series_number",
    "серия номер": "buyer.passport_series_number",
    "кем выдан покупателя": "buyer.passport_issued_by",
    "дата выдачи покупателя": "buyer.passport_issue_date",
    "код подразделения покупателя": "buyer.department_code",
    "прописка покупателя": "buyer.registration_address",
    "адрес покупателя": "buyer.registration_address",
    "автомобиль": "vehicle.make_model",
    "машина": "vehicle.make_model",
    "тип тс": "vehicle.vehicle_type",
    "тип": "vehicle.vehicle_type",
    "вин": "vehicle.vin",
    "vin": "vehicle.vin",
    "год": "vehicle.year",
    "цвет": "vehicle.color",
    "госномер": "vehicle.license_plate",
    "номер": "vehicle.license_plate",
    "птс": "vehicle.pts",
    "кем выдан птс": "vehicle.pts_issued_by",
    "дата выдачи птс": "vehicle.pts_issue_date",
    "дата птс": "vehicle.pts_issue_date",
    "стс": "vehicle.sts",
    "номер стс": "vehicle.sts",
    "кем выдан стс": "vehicle.sts_issued_by",
    "дата выдачи стс": "vehicle.sts_issue_date",
    "дата стс": "vehicle.sts_issue_date",
    "свидетельство": "vehicle.sts",
    "свидетельство о регистрации": "vehicle.sts",
    "свидетельство о регистрации тс": "vehicle.sts",
    "кузов": "vehicle.body_number",
    "двигатель": "vehicle.engine_number",
    "номер двигателя": "vehicle.engine_number",
    "мощность": "vehicle.engine_power",
    "мощность двигателя": "vehicle.engine_power",
    "объем": "vehicle.engine_volume",
    "объем двигателя": "vehicle.engine_volume",
    "рабочий объем": "vehicle.engine_volume",
    "шасси": "vehicle.chassis_number",
    "рама": "vehicle.chassis_number",
    "пробег": "vehicle.mileage",
}

MAIN_KEYBOARD = ReplyKeyboardMarkup(
    [
        ["📝 Новый договор", "🔍 Проверить данные"],
        ["✏️ Исправить поле", "📄 Создать договор"],
        ["💳 Проверить оплату"],
        ["⏭️ Пропустить фото", "❌ Отменить"],
        ["↩️ Шаг назад", "🧹 Очистить чат"],
        ["📊 Админ"],
    ],
    resize_keyboard=True,
)

PHONE_KEYBOARD = ReplyKeyboardMarkup(
    [[KeyboardButton("📱 Отправить номер телефона", request_contact=True)]],
    resize_keyboard=True,
    one_time_keyboard=True,
)

MONTH_NAMES = [
    "",
    "Январь",
    "Февраль",
    "Март",
    "Апрель",
    "Май",
    "Июнь",
    "Июль",
    "Август",
    "Сентябрь",
    "Октябрь",
    "Ноябрь",
    "Декабрь",
]

WEEKDAY_NAMES = ["Пн", "Вт", "Ср", "Чт", "Пт", "Сб", "Вс"]
PHOTO_TIP = (
    "Фото отправляй без бликов, не под углом, чтобы весь документ был в кадре. "
    "Текст и цифры должны быть четкими."
)

EDIT_GROUPS = {
    "seller_passport": {
        "title": "Паспорт продавца",
        "fields": [
            "seller.full_name",
            "seller.birth_date",
            "seller.passport_series_number",
            "seller.passport_issued_by",
            "seller.passport_issue_date",
            "seller.department_code",
        ],
    },
    "seller_registration": {
        "title": "Прописка продавца",
        "fields": ["seller.registration_address"],
    },
    "buyer_passport": {
        "title": "Паспорт покупателя",
        "fields": [
            "buyer.full_name",
            "buyer.birth_date",
            "buyer.passport_series_number",
            "buyer.passport_issued_by",
            "buyer.passport_issue_date",
            "buyer.department_code",
        ],
    },
    "buyer_registration": {
        "title": "Прописка покупателя",
        "fields": ["buyer.registration_address"],
    },
    "vehicle": {
        "title": "Автомобиль",
        "fields": [
            "vehicle.make_model",
            "vehicle.vin",
            "vehicle.vehicle_type",
            "vehicle.year",
            "vehicle.color",
            "vehicle.license_plate",
            "vehicle.mileage",
            "vehicle.engine_power",
            "vehicle.engine_volume",
            "vehicle.engine_number",
            "vehicle.chassis_number",
            "vehicle.body_number",
        ],
    },
    "vehicle_docs": {
        "title": "ПТС и СТС",
        "fields": [
            "vehicle.pts",
            "vehicle.pts_issued_by",
            "vehicle.pts_issue_date",
            "vehicle.sts",
            "vehicle.sts_issued_by",
            "vehicle.sts_issue_date",
        ],
    },
    "deal": {
        "title": "Город, дата, цена",
        "fields": ["deal.city", "deal.date", "deal.price"],
    },
}

STAGE_TO_SECTION = {
    "seller_passport": "seller",
    "seller_registration": "seller",
    "buyer_passport": "buyer",
    "buyer_registration": "buyer",
    "vehicle_sts_front": "vehicle",
    "vehicle_sts_back": "vehicle",
    "vehicle_pts": "vehicle",
}

TEMPLATE_VALUES = {
    "city": "deal.city",
    "contract_date": "deal.date",
    "price": "deal.price",
    "price_words": "deal.price_words",
    "seller_full_name": "seller.full_name",
    "seller_birth_date": "seller.birth_date",
    "seller_passport": "seller.passport_series_number",
    "seller_passport_issued_by": "seller.passport_issued_by",
    "seller_passport_issue_date": "seller.passport_issue_date",
    "seller_department_code": "seller.department_code",
    "seller_registration_address": "seller.registration_address",
    "buyer_full_name": "buyer.full_name",
    "buyer_birth_date": "buyer.birth_date",
    "buyer_passport": "buyer.passport_series_number",
    "buyer_passport_issued_by": "buyer.passport_issued_by",
    "buyer_passport_issue_date": "buyer.passport_issue_date",
    "buyer_department_code": "buyer.department_code",
    "buyer_registration_address": "buyer.registration_address",
    "vehicle_make_model": "vehicle.make_model",
    "vehicle_vin": "vehicle.vin",
    "vehicle_type": "vehicle.vehicle_type",
    "vehicle_year": "vehicle.year",
    "vehicle_color": "vehicle.color",
    "vehicle_license_plate": "vehicle.license_plate",
    "vehicle_pts": "vehicle.pts",
    "vehicle_pts_issued_by": "vehicle.pts_issued_by",
    "vehicle_pts_issue_date": "vehicle.pts_issue_date",
    "vehicle_sts": "vehicle.sts",
    "vehicle_sts_issued_by": "vehicle.sts_issued_by",
    "vehicle_sts_issue_date": "vehicle.sts_issue_date",
    "vehicle_body_number": "vehicle.body_number",
    "vehicle_engine_number": "vehicle.engine_number",
    "vehicle_engine_power": "vehicle.engine_power",
    "vehicle_engine_volume": "vehicle.engine_volume",
    "vehicle_chassis_number": "vehicle.chassis_number",
    "vehicle_mileage": "vehicle.mileage",
}


def ensure_json_file(path):
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.is_dir():
        raise RuntimeError(
            f"{path} is a directory, but a JSON file is required. "
            "Remove that directory on the server and restart the bot."
        )
    if not path.exists() or not path.read_text(encoding="utf-8").strip():
        path.write_text("{}", encoding="utf-8")


def load_json_file(path):
    ensure_json_file(path)
    try:
        with path.open("r", encoding="utf-8") as file:
            return json.load(file)
    except json.JSONDecodeError as error:
        raise RuntimeError(f"{path} contains invalid JSON. Replace it with {{}} and restart the bot.") from error


def ensure_storage():
    ensure_json_file(DATA_FILE)
    ensure_json_file(USERS_FILE)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    PHOTO_DIR.mkdir(parents=True, exist_ok=True)


def load_data():
    if not DATA_FILE.exists():
        return {}
    return load_json_file(DATA_FILE)


def save_data(data):
    with DATA_FILE.open("w", encoding="utf-8") as file:
        json.dump(data, file, ensure_ascii=False, indent=2)


def load_users():
    if not USERS_FILE.exists():
        return {}
    return load_json_file(USERS_FILE)


def save_users(users):
    with USERS_FILE.open("w", encoding="utf-8") as file:
        json.dump(users, file, ensure_ascii=False, indent=2)


def chat_id(update: Update):
    return str(update.effective_chat.id)


def user_id(update: Update):
    return str(update.effective_user.id)


def get_user(update: Update):
    return load_users().get(user_id(update))


def is_admin(update: Update):
    current_user_id = user_id(update)
    if current_user_id in ADMIN_IDS:
        return True
    user = get_user(update)
    return bool(user and user.get("role") == "admin")


def remaining_contracts(update: Update):
    if is_admin(update):
        return None
    user = get_user(update) or {}
    return max(0, FREE_CONTRACT_LIMIT - user.get("contracts_created", 0))


def register_user(update: Update, phone_number):
    users = load_users()
    current_user_id = user_id(update)
    is_first_user = not users and not ADMIN_IDS
    role = "admin" if is_first_user or current_user_id in ADMIN_IDS else "user"
    existing = users.get(current_user_id, {})
    users[current_user_id] = {
        **existing,
        "id": current_user_id,
        "chat_id": chat_id(update),
        "phone": phone_number,
        "first_name": update.effective_user.first_name or "",
        "last_name": update.effective_user.last_name or "",
        "username": update.effective_user.username or "",
        "registered_at": existing.get("registered_at") or datetime.now().strftime("%d.%m.%Y %H:%M"),
        "role": existing.get("role") or role,
        "deals_started": existing.get("deals_started", 0),
        "contracts_created": existing.get("contracts_created", 0),
    }
    save_users(users)
    return users[current_user_id]


def increment_user_stat(update: Update, field):
    users = load_users()
    current_user_id = user_id(update)
    if current_user_id not in users:
        return
    users[current_user_id][field] = users[current_user_id].get(field, 0) + 1
    save_users(users)


def payment_credentials_ready():
    return bool(YOOKASSA_SHOP_ID and YOOKASSA_SECRET_KEY)


def yookassa_request(method, path, payload=None, idempotence_key=None):
    url = f"https://api.yookassa.ru/v3{path}"
    data = None
    headers = {"Content-Type": "application/json"}
    if payload is not None:
        data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    if idempotence_key:
        headers["Idempotence-Key"] = idempotence_key

    credentials = f"{YOOKASSA_SHOP_ID}:{YOOKASSA_SECRET_KEY}".encode("utf-8")
    headers["Authorization"] = "Basic " + base64.b64encode(credentials).decode("ascii")

    request = urllib.request.Request(url, data=data, headers=headers, method=method)
    try:
        with urllib.request.urlopen(request, timeout=30) as response:
            return json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as error:
        details = error.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"ЮKassa вернула ошибку {error.code}: {details}") from error


def create_yookassa_payment(update: Update):
    payment_id = str(uuid.uuid4())
    payload = {
        "amount": {
            "value": f"{float(CONTRACT_PRICE_RUB):.2f}",
            "currency": "RUB",
        },
        "capture": True,
        "payment_method_data": {
            "type": "sbp",
        },
        "confirmation": {
            "type": "redirect",
            "return_url": YOOKASSA_RETURN_URL,
        },
        "description": f"Договор купли-продажи авто для пользователя {user_id(update)}",
        "metadata": {
            "telegram_user_id": user_id(update),
            "telegram_chat_id": chat_id(update),
            "local_payment_id": payment_id,
        },
    }
    return yookassa_request("POST", "/payments", payload, idempotence_key=payment_id)


def get_yookassa_payment(payment_id):
    return yookassa_request("GET", f"/payments/{payment_id}")


def require_registration_text():
    return "Для работы с ботом нужно зарегистрироваться. Нажми кнопку и отправь номер телефона."


def get_session(update: Update):
    data = load_data()
    return data.get(chat_id(update))


def save_session(update: Update, session):
    data = load_data()
    data[chat_id(update)] = session
    save_data(data)


def remember_message(update: Update, message):
    session = get_session(update)
    if not session or not message:
        return
    message_ids = session.setdefault("message_ids", [])
    if message.message_id not in message_ids:
        message_ids.append(message.message_id)
    session["message_ids"] = message_ids[-120:]
    save_session(update, session)


def remember_incoming_message(update: Update):
    if update.message:
        remember_message(update, update.message)


async def delete_message_safely(bot, chat, message_id):
    if not CLEAN_CHAT or not message_id:
        return
    try:
        await bot.delete_message(chat_id=chat, message_id=message_id)
    except Exception:
        pass


async def send_tracked_message(update: Update, context: ContextTypes.DEFAULT_TYPE, text, reply_markup=None):
    message = await context.bot.send_message(
        chat_id=update.effective_chat.id,
        text=text,
        reply_markup=reply_markup,
    )
    remember_message(update, message)
    return message


async def delete_tracked_messages(update: Update, context: ContextTypes.DEFAULT_TYPE):
    session = get_session(update)
    if not session:
        return
    message_ids = set(session.get("message_ids", []))
    if session.get("last_prompt_message_id"):
        message_ids.add(session["last_prompt_message_id"])
    for message_id in message_ids:
        await delete_message_safely(context.bot, update.effective_chat.id, message_id)
    session["message_ids"] = []
    session["last_prompt_message_id"] = None
    save_session(update, session)


async def delete_incoming_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    return


async def delete_last_prompt(update: Update, context: ContextTypes.DEFAULT_TYPE):
    return


async def send_clean_prompt(update: Update, context: ContextTypes.DEFAULT_TYPE, text):
    session = get_session(update)
    if session:
        await delete_message_safely(
            context.bot,
            update.effective_chat.id,
            session.get("last_prompt_message_id"),
        )
    message = await context.bot.send_message(
        chat_id=update.effective_chat.id,
        text=text,
        reply_markup=MAIN_KEYBOARD,
    )
    session = get_session(update)
    if session:
        session["last_prompt_message_id"] = message.message_id
        message_ids = session.setdefault("message_ids", [])
        if message.message_id not in message_ids:
            message_ids.append(message.message_id)
        session["message_ids"] = message_ids[-120:]
        save_session(update, session)
    return message


def merge_values(target, source):
    for key, value in source.items():
        if isinstance(value, dict) and isinstance(target.get(key), dict):
            merge_values(target[key], value)
        elif value not in (None, "", [], {}):
            target[key] = str(value).strip()


def normalize_date(value):
    if not value:
        return ""
    match = re.search(r"\b(\d{1,2})[./-](\d{1,2})[./-](\d{2,4})\b", str(value))
    if not match:
        return ""
    day, month, year = match.groups()
    if len(year) == 2:
        year = "20" + year if int(year) < 30 else "19" + year
    return f"{int(day):02d}.{int(month):02d}.{year}"


def parse_normalized_date(value):
    normalized = normalize_date(value)
    if not normalized:
        return None
    try:
        return datetime.strptime(normalized, "%d.%m.%Y")
    except ValueError:
        return None


def is_valid_passport_issue_date(issue_date, birth_date):
    issue = parse_normalized_date(issue_date)
    birth = parse_normalized_date(birth_date)
    if not issue:
        return False

    if issue > datetime.now():
        return False

    if birth:
        min_issue_year = birth.year + 14
        min_issue_date = birth.replace(year=min_issue_year)
        if issue < min_issue_date:
            return False

    return True


def clean_extracted_data(data):
    for person in ("seller", "buyer"):
        if person not in data or not isinstance(data[person], dict):
            continue
        data[person]["birth_date"] = normalize_date(data[person].get("birth_date", ""))
        data[person]["passport_issue_date"] = normalize_date(data[person].get("passport_issue_date", ""))
        if not is_valid_passport_issue_date(
            data[person].get("passport_issue_date", ""),
            data[person].get("birth_date", ""),
        ):
            data[person]["passport_issue_date"] = ""
    vehicle = data.get("vehicle")
    if isinstance(vehicle, dict):
        vehicle["pts_issue_date"] = normalize_date(vehicle.get("pts_issue_date", ""))
        vehicle["sts_issue_date"] = normalize_date(vehicle.get("sts_issue_date", ""))
    return data


def get_value(deal, field_path):
    section, field = field_path.split(".", 1)
    return deal.get(section, {}).get(field, "")


def set_value(deal, field_path, value):
    section, field = field_path.split(".", 1)
    if section not in deal or field not in deal[section]:
        return False
    deal[section][field] = value.strip()
    return True


def normalize_label(text):
    text = text.lower().replace("ё", "е")
    text = re.sub(r"[^a-zа-я0-9 ]+", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def normalize_button_text(text):
    return re.sub(r"[^a-zа-яё0-9 ]+", "", text.lower()).replace("ё", "е").strip()


def resolve_field(label):
    normalized = normalize_label(label)
    if normalized in FIELD_ALIASES:
        return FIELD_ALIASES[normalized]

    for alias, field in FIELD_ALIASES.items():
        if normalized.startswith(alias) or alias in normalized:
            return field
    return None


def parse_city_date(text):
    parts = [part.strip() for part in re.split(r"[,;\n]+", text) if part.strip()]
    city = parts[0] if parts else ""
    date = ""

    date_match = re.search(r"\d{1,2}[./-]\d{1,2}[./-]\d{2,4}", text)
    if date_match:
        date = date_match.group(0).replace("-", ".").replace("/", ".")
    elif len(parts) > 1:
        date = parts[1]

    return city, date


def build_calendar(year, month):
    rows = [
        [
            InlineKeyboardButton("<", callback_data=f"calendar:prev:{year}:{month}"),
            InlineKeyboardButton(f"{MONTH_NAMES[month]} {year}", callback_data="calendar:noop"),
            InlineKeyboardButton(">", callback_data=f"calendar:next:{year}:{month}"),
        ],
        [InlineKeyboardButton(day, callback_data="calendar:noop") for day in WEEKDAY_NAMES],
    ]

    for week in calendar.monthcalendar(year, month):
        row = []
        for day in week:
            if day == 0:
                row.append(InlineKeyboardButton(" ", callback_data="calendar:noop"))
            else:
                row.append(InlineKeyboardButton(str(day), callback_data=f"calendar:day:{year}:{month}:{day}"))
        rows.append(row)
    return InlineKeyboardMarkup(rows)


def build_edit_keyboard():
    rows = [
        [
            InlineKeyboardButton("🪪 Паспорт продавца", callback_data="edit:seller_passport"),
            InlineKeyboardButton("🏠 Прописка продавца", callback_data="edit:seller_registration"),
        ],
        [
            InlineKeyboardButton("🪪 Паспорт покупателя", callback_data="edit:buyer_passport"),
            InlineKeyboardButton("🏠 Прописка покупателя", callback_data="edit:buyer_registration"),
        ],
        [
            InlineKeyboardButton("🚗 Автомобиль", callback_data="edit:vehicle"),
            InlineKeyboardButton("📑 ПТС и СТС", callback_data="edit:vehicle_docs"),
        ],
        [InlineKeyboardButton("📍 Город, дата, цена", callback_data="edit:deal")],
    ]
    return InlineKeyboardMarkup(rows)


def build_all_fields_keyboard():
    rows = []
    row = []
    for field, label in FIELD_LABELS.items():
        row.append(InlineKeyboardButton(f"{field_icon(field)} {label}", callback_data=f"field:{field}"))
        if len(row) == 2:
            rows.append(row)
            row = []
    if row:
        rows.append(row)
    return InlineKeyboardMarkup(rows)


def field_icon(field):
    if field.startswith("seller.") or field.startswith("buyer."):
        if "registration" in field:
            return "🏠"
        if "passport" in field or "department" in field:
            return "🪪"
        return "👤"
    if field.startswith("vehicle."):
        if "pts" in field or "sts" in field:
            return "📑"
        if field in ("vehicle.vin", "vehicle.license_plate"):
            return "🔢"
        return "🚗"
    if field == "deal.price":
        return "💰"
    if field == "deal.date":
        return "📅"
    if field == "deal.city":
        return "📍"
    return "✏️"


def format_edit_group(deal, group_key):
    group = EDIT_GROUPS[group_key]
    lines = [
        group["title"],
        "",
        "Скопируй нужные строки, исправь и отправь обратно:",
        "",
    ]
    for field in group["fields"]:
        value = get_value(deal, field)
        lines.append(f"{FIELD_LABELS[field]} - {value}")
    return "\n".join(lines)


def format_edit_field(deal, field):
    label = FIELD_LABELS[field]
    value = get_value(deal, field)
    return f"{label} - {value}"


def shift_month(year, month, direction):
    month += direction
    if month < 1:
        return year - 1, 12
    if month > 12:
        return year + 1, 1
    return year, month


async def ask_contract_date(update: Update):
    today = datetime.now()
    await update.message.reply_text(
        "Выбери дату договора:",
        reply_markup=build_calendar(today.year, today.month),
    )


def parse_corrections(text):
    corrections = []
    chunks = [chunk.strip() for chunk in re.split(r"[\n,;]+", text) if chunk.strip()]

    for chunk in chunks:
        if ":" in chunk:
            label, value = chunk.split(":", 1)
        elif " - " in chunk:
            label, value = chunk.split(" - ", 1)
        elif "-" in chunk:
            label, value = chunk.split("-", 1)
        else:
            continue

        field = resolve_field(label)
        value = value.strip()
        if field and value:
            corrections.append((field, value))

    return corrections


def missing_fields(deal):
    required = [
        "seller.full_name",
        "seller.passport_series_number",
        "seller.registration_address",
        "buyer.full_name",
        "buyer.passport_series_number",
        "buyer.registration_address",
        "vehicle.make_model",
        "vehicle.vin",
        "vehicle.mileage",
        "deal.city",
        "deal.price",
    ]
    return [field for field in required if not get_value(deal, field)]


def format_deal(deal):
    lines = ["Проверь данные:"]
    for field, label in FIELD_LABELS.items():
        value = get_value(deal, field) or "не заполнено"
        mark = "✅" if get_value(deal, field) else "❗"
        lines.append(f"{mark} {label} - {value}")
    return "\n".join(lines)


def extract_json(text):
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?", "", text).strip()
        text = re.sub(r"```$", "", text).strip()
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if not match:
        return {}
    return json.loads(match.group(0))


def photo_to_data_url(photo_path):
    encoded = base64.b64encode(photo_path.read_bytes()).decode("utf-8")
    return f"data:image/jpeg;base64,{encoded}"


def yandex_request(url, payload, headers):
    request = urllib.request.Request(
        url,
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers={"Content-Type": "application/json", **headers},
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=60) as response:
            return json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as error:
        details = error.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"Yandex API вернул ошибку {error.code}: {details}") from error
    except urllib.error.URLError as error:
        raise RuntimeError(f"Не удалось подключиться к Yandex API: {error}") from error


def yandex_ocr_text(photo_path):
    if not YANDEX_API_KEY:
        raise RuntimeError("Добавь YANDEX_API_KEY в файл .env")

    payload = {
        "content": base64.b64encode(photo_path.read_bytes()).decode("utf-8"),
        "mimeType": "image/jpeg",
        "languageCodes": ["ru", "en"],
        "model": YANDEX_OCR_MODEL,
    }
    result = yandex_request(
        "https://ocr.api.cloud.yandex.net/ocr/v1/recognizeText",
        payload,
        {"Authorization": f"Api-Key {YANDEX_API_KEY}"},
    )
    text = result.get("textAnnotation", {}).get("fullText", "").strip()
    if not text:
        raise RuntimeError("Yandex OCR не распознал текст на фото")
    return text


def yandex_extract_data(prompt, recognized_text):
    if not YANDEX_API_KEY:
        raise RuntimeError("Добавь YANDEX_API_KEY в файл .env")
    if not YANDEX_FOLDER_ID:
        raise RuntimeError("Добавь YANDEX_FOLDER_ID в файл .env")

    payload = {
        "model": f"gpt://{YANDEX_FOLDER_ID}/{YANDEX_GPT_MODEL}",
        "messages": [
            {
                "role": "user",
                "content": (
                    f"{prompt}\n\n"
                    "Ниже текст, распознанный OCR с фото документа. "
                    "Заполняй JSON только по этому тексту, не придумывай данные.\n\n"
                    f"{recognized_text}"
                ),
            }
        ],
        "temperature": 0,
        "max_tokens": 2000,
        "stream": False,
    }
    result = yandex_request(
        "https://llm.api.cloud.yandex.net/v1/chat/completions",
        payload,
        {"Authorization": f"Api-Key {YANDEX_API_KEY}", "OpenAI-Project": YANDEX_FOLDER_ID},
    )
    content = result["choices"][0]["message"]["content"]
    return clean_extracted_data(extract_json(content))


def recognize_document(photo_path, stage):
    if AI_PROVIDER == "openai" and not OPENAI_API_KEY:
        raise RuntimeError("Добавь OPENAI_API_KEY в файл .env")

    section = STAGE_TO_SECTION[stage]
    prompt = f"""
Ты внимательно переписываешь данные с фото документа для черновика договора купли-продажи автомобиля.
Верни только JSON без пояснений.
Документ: {stage}. Заполняй только раздел {section}.

Правила:
- Не придумывай данные и не восстанавливай их по памяти.
- Если поле не видно, размыто, закрыто пальцем или ты не уверен, оставь пустую строку.
- Для ФИО в российском паспорте используй только видимые строки "Фамилия", "Имя", "Отчество".
- Запрещено подставлять похожее, вероятное или распространенное ФИО.
- Если хотя бы фамилия, имя или отчество читаются неуверенно, оставь full_name пустым.
- Не используй данные из других документов, прошлых запросов, примеров или памяти.
- Если на фото мужской паспорт, не возвращай женское ФИО; если пол и ФИО противоречат друг другу, оставь full_name пустым.
- Все даты возвращай только в формате ДД.ММ.ГГГГ.
- Не пиши даты словами, например "15 декабря 2002".
- Для российского паспорта дата выдачи находится рядом с подписью "Дата выдачи" на верхней странице паспорта.
- Дату выдачи паспорта можно брать только если она видна цифрами на фото, например 29.09.2017.
- Если дата выдачи не читается полностью именно цифрами, верни пустую строку.
- Поле passport_issued_by бери из блока "Паспорт выдан" на верхней странице паспорта.
- "Паспорт выдан" часто занимает 1-3 строки. Объедини все видимые строки этого блока в одну строку.
- Блок "Паспорт выдан" обычно находится выше даты выдачи и кода подразделения.
- Если видна хотя бы часть органа выдачи паспорта, запиши ее в passport_issued_by; не оставляй поле пустым при читаемом тексте.
- Не путай орган выдачи паспорта с местом рождения, адресом регистрации, печатью или машинно-считываемой строкой.
- Не путай дату выдачи с датой рождения, кодом подразделения, номером паспорта, машинно-считываемой строкой или красными вертикальными цифрами серии/номера.
- Серия и номер паспорта должны быть только из видимых цифр на странице паспорта.
- Адрес регистрации заполняй только по фото прописки. Если сейчас фото разворота паспорта с фото, адрес регистрации оставь пустым.
- ФИО собирай из фамилии, имени и отчества как одну строку.
- Если это СТС, переписывай только данные, которые реально видны на отправленной стороне СТС.
- Для СТС первая сторона обычно содержит марку/модель, VIN, год, цвет и госномер.
- Для СТС вторая сторона обычно содержит серию и номер свидетельства о регистрации ТС. Это поле записывай в sts.
- Номер СТС может быть подписан как "Свидетельство о регистрации ТС", "Регистрационный документ", "серия", "номер" или быть напечатан отдельной крупной строкой.
- Если видишь серию и номер СТС, объедини их в одну строку, например "99 99 123456" или "9999 123456", и запиши в sts.
- Не путай номер СТС с госномером, VIN, ПТС, номером кузова, номером двигателя или кодом подразделения.
- Если это вторая сторона СТС и номер СТС виден четко, поле sts обязательно должно быть заполнено.
- Если на СТС видны кем выдано свидетельство и дата выдачи, запиши их в sts_issued_by и sts_issue_date.
- Если кем выдано СТС или дата выдачи СТС не видны, оставь соответствующее поле пустым.
- Если видны тип ТС, мощность двигателя в л.с. или кВт, рабочий объем в куб. см, номер двигателя, номер шасси/рамы или номер кузова, перепиши их.
- Если это ПТС, переписывай данные автомобиля, тип ТС, мощность, рабочий объем и номер ПТС только с фото ПТС.
- В ПТС строка 23 "Наименование организации, выдавшей паспорт" - это pts_issued_by.
- В ПТС строка 25 "Дата выдачи паспорта" - это pts_issue_date.
- Дата выдачи ПТС часто закрыта печатью. Если она не видна полностью цифрами, оставь pts_issue_date пустым.
- Пробег обычно не указан в паспорте/СТС. Не придумывай пробег, оставь пустым, если он не виден.

Схема JSON:
{{
  "seller": {{
    "full_name": "", "birth_date": "", "passport_series_number": "",
    "passport_issued_by": "", "passport_issue_date": "", "department_code": "",
    "registration_address": ""
  }},
  "buyer": {{
    "full_name": "", "birth_date": "", "passport_series_number": "",
    "passport_issued_by": "", "passport_issue_date": "", "department_code": "",
    "registration_address": ""
  }},
  "vehicle": {{
    "make_model": "", "vin": "", "vehicle_type": "", "year": "", "color": "",
    "license_plate": "", "pts": "", "pts_issued_by": "", "pts_issue_date": "",
    "sts": "", "sts_issued_by": "", "sts_issue_date": "",
    "body_number": "", "engine_number": "",
    "engine_power": "", "engine_volume": "", "chassis_number": "",
    "mileage": ""
  }}
}}
""".strip()

    if AI_PROVIDER == "yandex":
        recognized_text = yandex_ocr_text(photo_path)
        return yandex_extract_data(prompt, recognized_text)
    if AI_PROVIDER != "openai":
        raise RuntimeError("AI_PROVIDER должен быть openai или yandex")

    client = OpenAI(api_key=OPENAI_API_KEY)
    request = {
        "model": OPENAI_MODEL,
        "input": [
            {
                "role": "user",
                "content": [
                    {"type": "input_text", "text": prompt},
                    {"type": "input_image", "image_url": photo_to_data_url(photo_path), "detail": "high"},
                ],
            }
        ],
    }
    if not OPENAI_MODEL.startswith("gpt-5"):
        request["temperature"] = 0

    response = client.responses.create(**request)
    return clean_extracted_data(extract_json(response.output_text))


def add_paragraph(document, text="", bold=False):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return paragraph


def price_to_words(price):
    digits = re.sub(r"\D+", "", str(price))
    if not digits:
        return ""

    number = int(digits)
    if number == 0:
        return "\u043d\u043e\u043b\u044c"

    units = [
        ["", "\u043e\u0434\u0438\u043d", "\u0434\u0432\u0430", "\u0442\u0440\u0438", "\u0447\u0435\u0442\u044b\u0440\u0435", "\u043f\u044f\u0442\u044c", "\u0448\u0435\u0441\u0442\u044c", "\u0441\u0435\u043c\u044c", "\u0432\u043e\u0441\u0435\u043c\u044c", "\u0434\u0435\u0432\u044f\u0442\u044c"],
        ["", "\u043e\u0434\u043d\u0430", "\u0434\u0432\u0435", "\u0442\u0440\u0438", "\u0447\u0435\u0442\u044b\u0440\u0435", "\u043f\u044f\u0442\u044c", "\u0448\u0435\u0441\u0442\u044c", "\u0441\u0435\u043c\u044c", "\u0432\u043e\u0441\u0435\u043c\u044c", "\u0434\u0435\u0432\u044f\u0442\u044c"],
    ]
    teens = {
        10: "\u0434\u0435\u0441\u044f\u0442\u044c",
        11: "\u043e\u0434\u0438\u043d\u043d\u0430\u0434\u0446\u0430\u0442\u044c",
        12: "\u0434\u0432\u0435\u043d\u0430\u0434\u0446\u0430\u0442\u044c",
        13: "\u0442\u0440\u0438\u043d\u0430\u0434\u0446\u0430\u0442\u044c",
        14: "\u0447\u0435\u0442\u044b\u0440\u043d\u0430\u0434\u0446\u0430\u0442\u044c",
        15: "\u043f\u044f\u0442\u043d\u0430\u0434\u0446\u0430\u0442\u044c",
        16: "\u0448\u0435\u0441\u0442\u043d\u0430\u0434\u0446\u0430\u0442\u044c",
        17: "\u0441\u0435\u043c\u043d\u0430\u0434\u0446\u0430\u0442\u044c",
        18: "\u0432\u043e\u0441\u0435\u043c\u043d\u0430\u0434\u0446\u0430\u0442\u044c",
        19: "\u0434\u0435\u0432\u044f\u0442\u043d\u0430\u0434\u0446\u0430\u0442\u044c",
    }
    tens = ["", "", "\u0434\u0432\u0430\u0434\u0446\u0430\u0442\u044c", "\u0442\u0440\u0438\u0434\u0446\u0430\u0442\u044c", "\u0441\u043e\u0440\u043e\u043a", "\u043f\u044f\u0442\u044c\u0434\u0435\u0441\u044f\u0442", "\u0448\u0435\u0441\u0442\u044c\u0434\u0435\u0441\u044f\u0442", "\u0441\u0435\u043c\u044c\u0434\u0435\u0441\u044f\u0442", "\u0432\u043e\u0441\u0435\u043c\u044c\u0434\u0435\u0441\u044f\u0442", "\u0434\u0435\u0432\u044f\u043d\u043e\u0441\u0442\u043e"]
    hundreds = ["", "\u0441\u0442\u043e", "\u0434\u0432\u0435\u0441\u0442\u0438", "\u0442\u0440\u0438\u0441\u0442\u0430", "\u0447\u0435\u0442\u044b\u0440\u0435\u0441\u0442\u0430", "\u043f\u044f\u0442\u044c\u0441\u043e\u0442", "\u0448\u0435\u0441\u0442\u044c\u0441\u043e\u0442", "\u0441\u0435\u043c\u044c\u0441\u043e\u0442", "\u0432\u043e\u0441\u0435\u043c\u044c\u0441\u043e\u0442", "\u0434\u0435\u0432\u044f\u0442\u044c\u0441\u043e\u0442"]
    groups = [
        ("", "", "", 0),
        ("\u0442\u044b\u0441\u044f\u0447\u0430", "\u0442\u044b\u0441\u044f\u0447\u0438", "\u0442\u044b\u0441\u044f\u0447", 1),
        ("\u043c\u0438\u043b\u043b\u0438\u043e\u043d", "\u043c\u0438\u043b\u043b\u0438\u043e\u043d\u0430", "\u043c\u0438\u043b\u043b\u0438\u043e\u043d\u043e\u0432", 0),
        ("\u043c\u0438\u043b\u043b\u0438\u0430\u0440\u0434", "\u043c\u0438\u043b\u043b\u0438\u0430\u0440\u0434\u0430", "\u043c\u0438\u043b\u043b\u0438\u0430\u0440\u0434\u043e\u0432", 0),
    ]

    def group_word(value, forms, gender):
        parts = []
        h = value // 100
        t = (value % 100) // 10
        u = value % 10
        if h:
            parts.append(hundreds[h])
        if t == 1:
            parts.append(teens[t * 10 + u])
        else:
            if t:
                parts.append(tens[t])
            if u:
                parts.append(units[gender][u])
        if forms[0]:
            last_two = value % 100
            last = value % 10
            if 11 <= last_two <= 19:
                parts.append(forms[2])
            elif last == 1:
                parts.append(forms[0])
            elif 2 <= last <= 4:
                parts.append(forms[1])
            else:
                parts.append(forms[2])
        return parts

    parts = []
    group_index = 0
    while number:
        value = number % 1000
        if value:
            name1, name2, name5, gender = groups[group_index]
            parts = group_word(value, (name1, name2, name5), gender) + parts
        number //= 1000
        group_index += 1

    return " ".join(parts)


def build_template_replacements(deal):
    deal.setdefault("deal", {})
    deal["deal"]["price_words"] = price_to_words(get_value(deal, "deal.price"))
    replacements = {
        "{{" + placeholder + "}}": get_value(deal, field_path)
        for placeholder, field_path in TEMPLATE_VALUES.items()
    }
    return replacements


def replace_in_paragraph(paragraph, replacements):
    original = "".join(run.text for run in paragraph.runs)
    replaced = original
    for placeholder, value in replacements.items():
        replaced = replaced.replace(placeholder, value or "")

    if replaced == original:
        return

    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = replaced
    else:
        paragraph.add_run(replaced)


def fill_template_document(document, replacements):
    for paragraph in document.paragraphs:
        replace_in_paragraph(paragraph, replacements)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)


def create_contract(deal, chat):
    OUTPUT_DIR.mkdir(exist_ok=True)
    filename = OUTPUT_DIR / f"dogovor_kuply_prodazhi_{chat}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    seller = deal["seller"]
    buyer = deal["buyer"]
    vehicle = deal["vehicle"]
    deal_info = deal["deal"]

    if TEMPLATE_FILE.exists():
        document = Document(TEMPLATE_FILE)
        fill_template_document(document, build_template_replacements(deal))
        document.save(filename)
        return filename

    document = Document()
    document.add_heading("Договор купли-продажи транспортного средства", level=1)
    add_paragraph(document, f"г. {deal_info['city']}                                           {deal_info['date']}")
    add_paragraph(document)
    add_paragraph(
        document,
        f"{seller['full_name']}, паспорт {seller['passport_series_number']}, выдан {seller['passport_issued_by']} "
        f"{seller['passport_issue_date']}, код подразделения {seller['department_code']}, зарегистрированный(ая) по адресу: "
        f"{seller['registration_address']}, именуемый(ая) далее 'Продавец', с одной стороны, и {buyer['full_name']}, "
        f"паспорт {buyer['passport_series_number']}, выдан {buyer['passport_issued_by']} {buyer['passport_issue_date']}, "
        f"код подразделения {buyer['department_code']}, зарегистрированный(ая) по адресу: {buyer['registration_address']}, "
        f"именуемый(ая) далее 'Покупатель', с другой стороны, заключили настоящий договор о нижеследующем:"
    )
    add_paragraph(document, "1. Предмет договора", bold=True)
    add_paragraph(
        document,
        f"Продавец продает, а Покупатель покупает транспортное средство: {vehicle['make_model']}, "
        f"год выпуска {vehicle['year']}, VIN {vehicle['vin']}, цвет {vehicle['color']}, "
        f"государственный регистрационный знак {vehicle['license_plate']}, ПТС {vehicle['pts']}, "
        f"СТС {vehicle['sts']}, номер кузова {vehicle['body_number']}, номер двигателя {vehicle['engine_number']}."
    )
    add_paragraph(document, "2. Цена договора", bold=True)
    add_paragraph(document, f"Стоимость транспортного средства составляет {deal_info['price']} рублей.")
    add_paragraph(document, "3. Передача транспортного средства", bold=True)
    add_paragraph(
        document,
        "Продавец подтверждает получение денежных средств, а Покупатель подтверждает получение транспортного средства, "
        "ключей и документов. Стороны претензий друг к другу не имеют."
    )
    add_paragraph(document, "4. Подписи сторон", bold=True)
    add_paragraph(document)
    add_paragraph(document, f"Продавец: __________________ / {seller['full_name']} /")
    add_paragraph(document)
    add_paragraph(document, f"Покупатель: ________________ / {buyer['full_name']} /")
    add_paragraph(document)
    add_paragraph(document, "Перед подписанием обязательно проверьте все данные и актуальные требования закона.")

    document.save(filename)
    return filename


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await delete_incoming_message(update, context)
    if not get_user(update):
        await update.message.reply_text(require_registration_text(), reply_markup=PHONE_KEYBOARD)
        return
    await update.message.reply_text(
        "Я помогу собрать договор купли-продажи авто.",
        reply_markup=MAIN_KEYBOARD,
    )


async def handle_contact(update: Update, context: ContextTypes.DEFAULT_TYPE):
    contact = update.message.contact
    if not contact:
        return
    if str(contact.user_id) != user_id(update):
        await update.message.reply_text("Нужно отправить свой номер через кнопку.", reply_markup=PHONE_KEYBOARD)
        return

    user = register_user(update, contact.phone_number)
    if user["role"] == "admin":
        text = "Регистрация готова. Ты админ бота.\n\n" + TEST_MODE_NOTICE
    else:
        text = "Регистрация готова. Можно создавать договор.\n\n" + TEST_MODE_NOTICE
    await update.message.reply_text(text, reply_markup=MAIN_KEYBOARD)


async def new_deal(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await delete_incoming_message(update, context)
    if not get_user(update):
        await update.message.reply_text(require_registration_text(), reply_markup=PHONE_KEYBOARD)
        return
    session = {
        "deal_id": str(uuid.uuid4()),
        "stage_index": 0,
        "deal": deepcopy(EMPTY_DEAL),
        "mode": "city",
        "pending": 0,
        "message_ids": [],
    }
    save_session(update, session)
    increment_user_stat(update, "deals_started")
    remember_incoming_message(update)
    await send_tracked_message(
        update,
        context,
        "Напиши город договора.\nНапример: Москва",
        reply_markup=MAIN_KEYBOARD,
    )


async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await delete_incoming_message(update, context)
    data = load_data()
    data.pop(chat_id(update), None)
    save_data(data)
    await update.message.reply_text("Текущий договор отменен.", reply_markup=MAIN_KEYBOARD)


async def clear_chat(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await delete_incoming_message(update, context)
    session = get_session(update)
    if not session:
        return

    message_ids = set(session.get("message_ids", []))
    if session.get("last_prompt_message_id"):
        message_ids.add(session["last_prompt_message_id"])

    for message_id in message_ids:
        await delete_message_safely(context.bot, update.effective_chat.id, message_id)

    session["message_ids"] = []
    session["last_prompt_message_id"] = None
    save_session(update, session)
    await send_tracked_message(update, context, "Чат очищен.", reply_markup=MAIN_KEYBOARD)


async def admin_report(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_admin(update):
        await send_tracked_message(update, context, "У тебя нет доступа к админ-отчету.", reply_markup=MAIN_KEYBOARD)
        return

    users = load_users()
    sessions = load_data()
    total_users = len(users)
    total_deals = sum(user.get("deals_started", 0) for user in users.values())
    total_contracts = sum(user.get("contracts_created", 0) for user in users.values())
    admins = sum(1 for user in users.values() if user.get("role") == "admin")
    limited_users = sum(1 for user in users.values() if user.get("role") != "admin" and user.get("contracts_created", 0) >= FREE_CONTRACT_LIMIT)

    lines = [
        "📊 Админ-отчет",
        "",
        f"Пользователей: {total_users}",
        f"Админов: {admins}",
        f"Начали договоров: {total_deals}",
        f"Создали договоров: {total_contracts}",
        f"Исчерпали тестовый лимит: {limited_users}",
        f"Активных сессий: {len(sessions)}",
        "",
        "Последние пользователи:",
    ]

    recent_users = sorted(users.values(), key=lambda user: user.get("registered_at", ""), reverse=True)[:10]
    for user in recent_users:
        name = " ".join(part for part in [user.get("first_name", ""), user.get("last_name", "")] if part).strip()
        username = f"@{user['username']}" if user.get("username") else "без username"
        lines.append(
            f"- {name or 'без имени'} | {username} | {user.get('phone', 'без телефона')} | "
            f"договоров: {user.get('contracts_created', 0)}"
        )

    await send_tracked_message(update, context, "\n".join(lines), reply_markup=MAIN_KEYBOARD)


async def step_back(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await delete_incoming_message(update, context)
    await delete_last_prompt(update, context)
    session = get_session(update)
    if not session:
        await send_tracked_message(update, context, "Сначала нажми «Новый договор».", reply_markup=MAIN_KEYBOARD)
        return

    mode = session.get("mode")
    if mode == "city":
        await send_tracked_message(update, context, "Ты уже в начале. Напиши город договора.", reply_markup=MAIN_KEYBOARD)
        return

    if mode == "date":
        session["mode"] = "city"
        save_session(update, session)
        await send_tracked_message(update, context, "Напиши город договора.\nНапример: Москва", reply_markup=MAIN_KEYBOARD)
        return

    if mode == "photos":
        session["stage_index"] = max(0, session.get("stage_index", 0) - 1)
        save_session(update, session)
        await send_next_step(update, session, context)
        return

    if mode == "price":
        session["mode"] = "photos"
        session["stage_index"] = max(0, min(len(STAGES) - 1, session.get("stage_index", len(STAGES)) - 1))
        save_session(update, session)
        await send_next_step(update, session, context)
        return

    if mode == "mileage":
        session["mode"] = "price"
        save_session(update, session)
        await send_clean_prompt(update, context, "Теперь напиши цену автомобиля.\nНапример: Цена - 850000")
        return

    session["mode"] = "mileage"
    save_session(update, session)
    await send_clean_prompt(update, context, "Напиши пробег.\nНапример: Пробег - 125000")


async def draft(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await delete_incoming_message(update, context)
    session = get_session(update)
    if not session:
        await update.message.reply_text("Сначала нажми «Новый договор».", reply_markup=MAIN_KEYBOARD)
        return
    text = format_deal(session["deal"])
    if session.get("pending", 0):
        text += "\n\nНекоторые фото еще обрабатываются. Проверь еще раз через минуту."
    await send_tracked_message(update, context, text[:3900], reply_markup=MAIN_KEYBOARD)
    await send_tracked_message(update, context, "Исправить по разделам:", reply_markup=build_edit_keyboard())
    await send_tracked_message(update, context, "Или выбери конкретное поле:", reply_markup=build_all_fields_keyboard())


async def set_field(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await delete_incoming_message(update, context)
    session = get_session(update)
    if not session:
        await update.message.reply_text("Сначала нажми «Новый договор».", reply_markup=MAIN_KEYBOARD)
        return
    if len(context.args) < 2:
        session["mode"] = "edit"
        save_session(update, session)
        await update.message.reply_text(
            "Напиши, что исправить.\nНапример: Покупатель - Анатолий\nМожно несколько строк сразу.",
            reply_markup=MAIN_KEYBOARD,
        )
        return

    field = context.args[0]
    value = " ".join(context.args[1:])
    if not set_value(session["deal"], field, value):
        await update.message.reply_text("Не понял поле. Нажми «Исправить поле» и напиши по-русски.", reply_markup=MAIN_KEYBOARD)
        return

    save_session(update, session)
    await update.message.reply_text("Записал.", reply_markup=MAIN_KEYBOARD)


async def skip(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await delete_incoming_message(update, context)
    session = get_session(update)
    if not session:
        await update.message.reply_text("Сначала нажми «Новый договор».", reply_markup=MAIN_KEYBOARD)
        return
    session["stage_index"] += 1
    if session["stage_index"] >= len(STAGES):
        session["mode"] = "price"
    save_session(update, session)
    await send_next_step(update, session, context)


async def confirm(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await delete_incoming_message(update, context)
    session = get_session(update)
    if not session:
        await update.message.reply_text("Сначала нажми «Новый договор».", reply_markup=MAIN_KEYBOARD)
        return

    remaining = remaining_contracts(update)
    if remaining == 0:
        await update.message.reply_text(
            "Лимит тестовых генераций закончился. В тестовом режиме доступно только 2 договора.",
            reply_markup=MAIN_KEYBOARD,
        )
        return

    if session.get("pending", 0):
        await update.message.reply_text("Фото еще обрабатываются. Подожди немного и нажми «Создать договор» еще раз.")
        return

    missing = missing_fields(session["deal"])
    if missing:
        lines = ["Не хватает данных. Нажми «Исправить поле» и допиши:"]
        lines.extend(FIELD_LABELS[field] for field in missing)
        await update.message.reply_text("\n".join(lines), reply_markup=MAIN_KEYBOARD)
        return

    if PAYMENTS_ENABLED:
        await request_payment(update, context)
        return

    await send_contract(update, context)


async def send_contract(update: Update, context: ContextTypes.DEFAULT_TYPE):
    session = get_session(update)
    if not session:
        await send_tracked_message(update, context, "Сначала нажми «📝 Новый договор».", reply_markup=MAIN_KEYBOARD)
        return

    remaining = remaining_contracts(update)
    if remaining == 0:
        await send_tracked_message(
            update,
            context,
            "Лимит тестовых генераций закончился. В тестовом режиме доступно только 2 договора.",
            reply_markup=MAIN_KEYBOARD,
        )
        return

    path = create_contract(session["deal"], chat_id(update))
    document_message = await update.message.reply_document(document=path.open("rb"), filename=path.name)
    increment_user_stat(update, "contracts_created")
    await delete_tracked_messages(update, context)
    session = get_session(update)
    if session:
        session["message_ids"] = [document_message.message_id]
        save_session(update, session)


async def request_payment(update: Update, context: ContextTypes.DEFAULT_TYPE):
    session = get_session(update)
    if not session:
        await send_tracked_message(update, context, "Сначала нажми «📝 Новый договор».", reply_markup=MAIN_KEYBOARD)
        return

    if session.get("payment", {}).get("status") == "succeeded":
        await send_contract(update, context)
        return

    if not payment_credentials_ready():
        await send_tracked_message(
            update,
            context,
            "Оплата ЮKassa еще не настроена. Добавь YOOKASSA_SHOP_ID и YOOKASSA_SECRET_KEY в .env.",
            reply_markup=MAIN_KEYBOARD,
        )
        return

    payment = session.get("payment")
    if not payment or payment.get("status") not in ("pending", "waiting_for_capture"):
        try:
            created = create_yookassa_payment(update)
        except Exception as error:
            await send_tracked_message(update, context, f"Не получилось создать платеж: {error}", reply_markup=MAIN_KEYBOARD)
            return

        confirmation_url = created.get("confirmation", {}).get("confirmation_url")
        payment = {
            "id": created.get("id"),
            "status": created.get("status"),
            "confirmation_url": confirmation_url,
            "created_at": datetime.now().strftime("%d.%m.%Y %H:%M"),
        }
        session["payment"] = payment
        save_session(update, session)

    keyboard = InlineKeyboardMarkup([
        [InlineKeyboardButton("💳 Оплатить через СБП", url=payment["confirmation_url"])],
    ])
    await send_tracked_message(
        update,
        context,
        f"Для создания договора нужно оплатить {CONTRACT_PRICE_RUB} руб. через СБП.\nПосле оплаты нажми «💳 Проверить оплату».",
        reply_markup=keyboard,
    )


async def check_payment(update: Update, context: ContextTypes.DEFAULT_TYPE):
    session = get_session(update)
    if not session:
        await send_tracked_message(update, context, "Сначала нажми «📝 Новый договор».", reply_markup=MAIN_KEYBOARD)
        return

    payment = session.get("payment")
    if not payment or not payment.get("id"):
        await send_tracked_message(update, context, "Платеж еще не создан. Нажми «📄 Создать договор».", reply_markup=MAIN_KEYBOARD)
        return

    if not payment_credentials_ready():
        await send_tracked_message(update, context, "Оплата ЮKassa еще не настроена в .env.", reply_markup=MAIN_KEYBOARD)
        return

    try:
        actual = get_yookassa_payment(payment["id"])
    except Exception as error:
        await send_tracked_message(update, context, f"Не получилось проверить оплату: {error}", reply_markup=MAIN_KEYBOARD)
        return

    status = actual.get("status")
    paid = actual.get("paid") is True or status == "succeeded"
    session["payment"]["status"] = status
    session["payment"]["paid"] = paid
    save_session(update, session)

    if paid:
        await send_contract(update, context)
        return

    await send_tracked_message(update, context, f"Оплата пока не прошла. Статус: {status}", reply_markup=MAIN_KEYBOARD)


async def send_next_step(update: Update, session, context=None):
    index = session["stage_index"]
    if index < len(STAGES):
        if context:
            await send_clean_prompt(update, context, STAGES[index][1])
        else:
            await update.message.reply_text(STAGES[index][1], reply_markup=MAIN_KEYBOARD)
        return
    text = "Теперь напиши цену автомобиля.\nНапример: Цена - 850000"
    if context:
        await send_clean_prompt(update, context, text)
    else:
        await update.message.reply_text(text, reply_markup=MAIN_KEYBOARD)


async def recognize_photo_background(bot, chat, deal_id, file_id, photo_path, stage):
    data = load_data()
    session = data.get(chat)
    if not session or session.get("deal_id") != deal_id:
        return

    try:
        PHOTO_DIR.mkdir(exist_ok=True)
        telegram_file = await bot.get_file(file_id)
        await telegram_file.download_to_drive(photo_path)
        extracted = recognize_document(photo_path, stage)
    except Exception:
        await bot.send_message(
            chat_id=chat,
            text="Одно фото не получилось прочитать. Потом можно дописать данные через «Исправить поле».",
            reply_markup=MAIN_KEYBOARD,
        )
        extracted = {}
    finally:
        try:
            photo_path.unlink()
        except OSError:
            pass

    data = load_data()
    session = data.get(chat)
    if not session or session.get("deal_id") != deal_id:
        return

    merge_values(session["deal"], extracted)
    session["pending"] = max(0, session.get("pending", 1) - 1)
    data[chat] = session
    save_data(data)


async def handle_photo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    remember_incoming_message(update)
    session = get_session(update)
    if not session:
        await update.message.reply_text("Сначала нажми «Новый договор».", reply_markup=MAIN_KEYBOARD)
        return

    index = session["stage_index"]
    if index >= len(STAGES):
        await update.message.reply_text("Все фото уже собраны. Нажми «Проверить данные».", reply_markup=MAIN_KEYBOARD)
        return

    stage, _ = STAGES[index]
    file_id = update.message.photo[-1].file_id
    message_id = update.message.message_id
    await delete_last_prompt(update, context)

    session["stage_index"] += 1
    session["pending"] = session.get("pending", 0) + 1
    if session["stage_index"] >= len(STAGES):
        session["mode"] = "price"
    save_session(update, session)

    photo_path = PHOTO_DIR / f"{chat_id(update)}_{stage}_{update.message.message_id}.jpg"
    context.application.create_task(
        recognize_photo_background(context.bot, chat_id(update), session["deal_id"], file_id, photo_path, stage)
    )
    await send_next_step(update, session, context)


async def handle_calendar(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    parts = query.data.split(":")
    action = parts[1]

    if action == "noop":
        return

    if action in ("prev", "next"):
        year = int(parts[2])
        month = int(parts[3])
        direction = -1 if action == "prev" else 1
        year, month = shift_month(year, month, direction)
        await query.edit_message_reply_markup(reply_markup=build_calendar(year, month))
        return

    if action == "day":
        year = int(parts[2])
        month = int(parts[3])
        day = int(parts[4])
        session = get_session(update)
        if not session:
            await query.edit_message_text("Сначала нажми «Новый договор».")
            return

        session["deal"]["deal"]["date"] = f"{day:02d}.{month:02d}.{year}"
        session["mode"] = "photos"
        save_session(update, session)
        await query.edit_message_text(f"Дата: {day:02d}.{month:02d}.{year}")
        tip_message = await context.bot.send_message(
            chat_id=query.message.chat_id,
            text=PHOTO_TIP,
            reply_markup=MAIN_KEYBOARD,
        )
        message = await context.bot.send_message(
            chat_id=query.message.chat_id,
            text=STAGES[session["stage_index"]][1],
            reply_markup=MAIN_KEYBOARD,
        )
        session["last_prompt_message_id"] = message.message_id
        message_ids = session.setdefault("message_ids", [])
        for sent_message in (tip_message, message):
            if sent_message.message_id not in message_ids:
                message_ids.append(sent_message.message_id)
        session["message_ids"] = message_ids[-120:]
        save_session(update, session)


async def handle_edit_button(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    group_key = query.data.split(":", 1)[1]
    session = get_session(update)
    if not session:
        await query.message.reply_text("Сначала нажми «Новый договор».", reply_markup=MAIN_KEYBOARD)
        return
    if group_key not in EDIT_GROUPS:
        await query.message.reply_text("Не понял, что исправить.", reply_markup=MAIN_KEYBOARD)
        return

    session["mode"] = "edit"
    save_session(update, session)
    await send_tracked_message(update, context, format_edit_group(session["deal"], group_key), reply_markup=MAIN_KEYBOARD)


async def handle_field_button(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    field = query.data.split(":", 1)[1]
    session = get_session(update)
    if not session:
        await query.message.reply_text("Сначала нажми «Новый договор».", reply_markup=MAIN_KEYBOARD)
        return
    if field not in FIELD_LABELS:
        await query.message.reply_text("Не понял поле.", reply_markup=MAIN_KEYBOARD)
        return

    session["mode"] = "edit"
    save_session(update, session)
    await send_tracked_message(update, context, "Исправь строку ниже и отправь обратно:", reply_markup=MAIN_KEYBOARD)
    await send_tracked_message(update, context, format_edit_field(session["deal"], field))


async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    remember_incoming_message(update)
    text = update.message.text.strip()
    lowered = normalize_button_text(text)
    await delete_last_prompt(update, context)
    await delete_incoming_message(update, context)

    if lowered == "новый договор":
        await new_deal(update, context)
        return
    if lowered == "проверить данные":
        await draft(update, context)
        return
    if lowered == "создать договор":
        await confirm(update, context)
        return
    if lowered == "проверить оплату":
        await check_payment(update, context)
        return
    if lowered == "пропустить фото":
        await skip(update, context)
        return
    if lowered == "отменить":
        await cancel(update, context)
        return
    if lowered == "шаг назад":
        await step_back(update, context)
        return
    if lowered == "очистить чат":
        await clear_chat(update, context)
        return
    if lowered == "админ":
        await admin_report(update, context)
        return
    if lowered == "исправить поле":
        session = get_session(update)
        if not session:
            await update.message.reply_text("Сначала нажми «Новый договор».", reply_markup=MAIN_KEYBOARD)
            return
        session["mode"] = "edit"
        save_session(update, session)
        await update.message.reply_text(
            "Напиши, что исправить.\n"
            "Например:\n"
            "Покупатель - Анатолий\n"
            "Паспорт покупателя - 4512 345678\n"
            "Цена - 850000",
            reply_markup=MAIN_KEYBOARD,
        )
        return

    session = get_session(update)
    if not session:
        if not get_user(update):
            await update.message.reply_text(require_registration_text(), reply_markup=PHONE_KEYBOARD)
        else:
            await update.message.reply_text("Нажми «Новый договор».", reply_markup=MAIN_KEYBOARD)
        return

    mode = session.get("mode")

    if mode == "city":
        session["deal"]["deal"]["city"] = text
        session["mode"] = "date"
        save_session(update, session)
        await ask_contract_date(update)
        return

    if mode == "date":
        date = normalize_date(text)
        if not date:
            await update.message.reply_text("Выбери дату в календаре или напиши ее так: 15.05.2026")
            return
        session["deal"]["deal"]["date"] = date
        session["mode"] = "photos"
        save_session(update, session)
        await update.message.reply_text(PHOTO_TIP, reply_markup=MAIN_KEYBOARD)
        await send_next_step(update, session, context)
        return

    if mode == "price" and not parse_corrections(text):
        session["deal"]["deal"]["price"] = text
        session["mode"] = "mileage"
        save_session(update, session)
        await update.message.reply_text("Записал цену. Теперь напиши пробег.\nНапример: Пробег - 125000", reply_markup=MAIN_KEYBOARD)
        return

    if mode == "mileage" and not parse_corrections(text):
        session["deal"]["vehicle"]["mileage"] = text
        session["mode"] = "edit"
        save_session(update, session)
        await update.message.reply_text("Записал пробег. Теперь нажми «Проверить данные».", reply_markup=MAIN_KEYBOARD)
        return

    corrections = parse_corrections(text)
    if corrections:
        for field, value in corrections:
            set_value(session["deal"], field, value)
        session["mode"] = "edit"
        save_session(update, session)
        await update.message.reply_text("Записал. Нажми «Проверить данные».", reply_markup=MAIN_KEYBOARD)
        return

    await update.message.reply_text("Не понял. Можно нажать кнопку ниже.", reply_markup=MAIN_KEYBOARD)


def main():
    if not TELEGRAM_BOT_TOKEN:
        raise RuntimeError("Добавь TELEGRAM_BOT_TOKEN в файл .env")

    ensure_storage()
    logger.info("Starting Telegram contract bot. Data directory: %s", DATA_DIR)

    app = Application.builder().token(TELEGRAM_BOT_TOKEN).build()
    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("newdeal", new_deal))
    app.add_handler(CommandHandler("draft", draft))
    app.add_handler(CommandHandler("set", set_field))
    app.add_handler(CommandHandler("skip", skip))
    app.add_handler(CommandHandler("confirm", confirm))
    app.add_handler(CommandHandler("cancel", cancel))
    app.add_handler(CallbackQueryHandler(handle_calendar, pattern=r"^calendar:"))
    app.add_handler(CallbackQueryHandler(handle_edit_button, pattern=r"^edit:"))
    app.add_handler(CallbackQueryHandler(handle_field_button, pattern=r"^field:"))
    app.add_handler(MessageHandler(filters.CONTACT, handle_contact))
    app.add_handler(MessageHandler(filters.PHOTO, handle_photo))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))
    app.run_polling()


if __name__ == "__main__":
    main()
